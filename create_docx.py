from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

# 创建文档
doc = Document()

# 设置标题
title = doc.add_heading('基于核方法的机器学习案例工作总结', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 添加基本信息
doc.add_paragraph('姓名：李天昊')
doc.add_paragraph('学号：U202342573')
doc.add_paragraph('日期：2026年7月9日')

# 摘要
doc.add_heading('摘要', level=1)
doc.add_paragraph(
    '本文基于两个机器学习实践案例展开工作总结：第一个案例使用支持向量机（SVM）的四种核函数对鸢尾花数据集进行分类并对比决策边界；'
    '第二个案例使用高斯过程回归（GPR）对CO₂浓度时间序列数据进行建模与预测。两个案例均属于核方法范畴，分别从分类与回归两个角度体现了核函数在机器学习中的重要作用。'
    '本文将从数据、方法、结果、分析和总结五个部分进行系统梳理。'
)

# 一、数据
doc.add_heading('一、数据', level=1)
doc.add_heading('1.1 SVM分类数据集：鸢尾花（Iris）数据集', level=2)
doc.add_paragraph(
    '鸢尾花数据集是机器学习领域最经典的数据集之一，共包含150个样本，分为3个类别：山鸢尾（Setosa）、变色鸢尾（Versicolour）和维吉尼亚鸢尾（Virginica），每类各50个样本。'
    '每个样本包含4个特征：花萼长度、花萼宽度、花瓣长度和花瓣宽度。本案例仅使用前两个特征（花萼长度与花萼宽度）作为输入，以方便可视化决策边界。'
    '数据经过StandardScaler标准化处理，按70%训练集、30%测试集的比例划分。'
)

doc.add_heading('1.2 GPR回归数据集：CO₂浓度时间序列数据', level=2)
doc.add_paragraph(
    'CO₂浓度数据模拟了夏威夷冒纳罗亚（Mauna Loa）观测站的月度观测记录，时间范围从1958年3月至1971年12月，共166个数据点。'
    '数据具有明显的趋势性（长期上升）和周期性（季节性波动）特征。数据按85%训练集、15%测试集的比例划分，训练集141个样本，测试集25个样本。'
    '在模型训练前，对目标变量进行了均值中心化处理。'
)

# 二、方法
doc.add_heading('二、方法', level=1)
doc.add_heading('2.1 SVM核函数对比方法', level=2)
doc.add_paragraph(
    '支持向量机（SVM）是一种基于最大间隔原则的分类算法。对于非线性可分问题，SVM通过核函数将原始特征映射到高维空间，再在高维空间中寻找最优超平面。'
    '本案例对比了四种核函数：线性核（Linear）、多项式核（Polynomial，degree=3）、RBF核（γ=0.5）和Sigmoid核。所有核函数的惩罚参数C均设为1.0，多分类策略采用one-vs-rest。'
)

doc.add_heading('2.2 高斯过程回归方法', level=2)
doc.add_paragraph(
    '高斯过程回归（GPR）是一种非参数化的贝叶斯回归方法，通过先验分布和似然函数对数据进行建模。本案例采用复合核函数：'
)
doc.add_paragraph('K = C₁² · RBF(ℓ₁) + C₂² · ExpSineSquared(ℓ₂, P) + WhiteKernel(σ²)')
doc.add_paragraph(
    '其中RBF核用于建模长期趋势，ExpSineSquared核用于捕捉年度周期性，WhiteKernel用于建模观测噪声。'
    '模型通过最大化对数边际似然自动优化核函数的超参数。'
)

# 三、结果
doc.add_heading('三、结果', level=1)
doc.add_heading('3.1 SVM分类结果', level=2)
doc.add_paragraph('四种核函数在鸢尾花测试集上的准确率如下表所示：')

# 创建表格
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '核函数'
hdr_cells[1].text = '测试集准确率'

rows = [
    ('线性核（Linear）', '0.8000'),
    ('多项式核（Poly, d=3）', '0.7556'),
    ('RBF核（γ=0.5）', '0.7333'),
    ('Sigmoid核（Sigmoid）', '0.7556')
]

for i, (kernel, acc) in enumerate(rows, 1):
    cells = table.rows[i].cells
    cells[0].text = kernel
    cells[1].text = acc

doc.add_paragraph('四种核函数对应的决策边界如下图所示：')

# 插入SVM图片
doc.add_picture('C:/xxq/output/svm_kernels_comparison.png', width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('图1：SVM四种核函数在鸢尾花数据集（前两维）上的决策边界对比').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('3.2 GPR预测结果', level=2)
doc.add_paragraph(
    '优化后的复合核函数参数为：11.6² · RBF(length_scale=20.0) + 5.0² · ExpSineSquared(length_scale=4.09, periodicity=1.0) + WhiteKernel(noise_level=0.0773)。'
    '优化后的对数边际似然为-44.10。模型在测试集上的MAE为0.317 ppm，RMSE为0.387 ppm。'
)

doc.add_paragraph('GPR对训练数据、测试数据以及未来至1978年的预测结果如下图所示：')

# 插入GPR图片
doc.add_picture('C:/xxq/output/gpr_co2_prediction.png', width=Inches(6.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('图2：高斯过程回归对CO₂浓度的建模与预测结果').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('1975—1978年的具体预测结果如下表所示：')

future_table = doc.add_table(rows=5, cols=3)
future_table.style = 'Light Grid Accent 1'
hdr_cells = future_table.rows[0].cells
hdr_cells[0].text = '年份'
hdr_cells[1].text = '预测CO₂浓度（ppm）'
hdr_cells[2].text = '95%置信区间'

future_rows = [
    ('1975', '330.55', '±1.22'),
    ('1976', '331.19', '±1.50'),
    ('1977', '331.76', '±1.83'),
    ('1978', '332.26', '±2.22')
]

for i, (year, pred, ci) in enumerate(future_rows, 1):
    cells = future_table.rows[i].cells
    cells[0].text = year
    cells[1].text = pred
    cells[2].text = ci

# 四、分析
doc.add_heading('四、分析', level=1)
doc.add_heading('4.1 SVM核函数选择分析', level=2)
doc.add_paragraph(
    '线性核在鸢尾花前两维特征上取得了最高准确率，这并不意味着数据整体线性可分，而是因为仅使用两个特征时数据分布恰好呈现较强的线性可分趋势。'
    'RBF核虽然具有强大的非线性建模能力，但准确率最低，说明当数据具有较好线性结构时，过于复杂的核函数容易导致过拟合，反而降低泛化能力。'
    '从决策边界来看，RBF核和Sigmoid核的边界过于复杂，而线性核和多项式核的边界相对平滑。这验证了模型复杂度与数据复杂度相匹配的原则。'
)

doc.add_heading('4.2 GPR复合核函数分析', level=2)
doc.add_paragraph(
    '复合核函数能够同时捕捉数据中的长期趋势、周期性规律和随机噪声。RBF长度尺度为20.0说明CO₂浓度长期趋势变化缓慢；'
    'ExpSineSquared周期固定为1.0，对应年度季节性波动。测试集上的MAE和RMSE分别为0.317 ppm和0.387 ppm，说明模型具有良好的预测精度。'
    '更重要的是，GPR提供了预测的不确定性信息：在训练数据覆盖范围内置信区间较窄，在外推区域置信区间逐渐扩大，这一特性对于时间序列预测具有重要意义。'
)

doc.add_heading('4.3 两个案例的共性认识', level=2)
doc.add_paragraph(
    'SVM和GPR虽然分别属于分类和回归任务，但二者都依赖于核函数这一核心概念。核函数的本质是隐式地度量样本之间的相似性，并通过特征空间的映射使原本复杂的问题变得可解。'
)

# 五、总结
doc.add_heading('五、总结', level=1)
doc.add_paragraph(
    '通过今天的两个案例实践，我对核方法在机器学习中的应用有了更加系统和深入的理解。在SVM分类实验中，我认识到核函数的选择应结合数据分布和任务需求；'
    '在GPR回归实验中，我体会到了复合核函数的灵活性和贝叶斯方法在不确定性量化方面的独特优势。'
    '今后的学习中，我将进一步尝试超参数调优、交叉验证和更复杂核结构的设计，以提升模型的泛化能力和预测精度。'
)

# 附录
doc.add_heading('附录', level=1)
doc.add_paragraph('代码来源：')
doc.add_paragraph('SVM案例：C:/xxq/project_four_A.py', style='List Bullet')
doc.add_paragraph('GPR案例：C:/xxq/project_three_part_4_A.py', style='List Bullet')
doc.add_paragraph('结果文件：')
doc.add_paragraph('SVM对比图：C:/xxq/output/svm_kernels_comparison.png', style='List Bullet')
doc.add_paragraph('GPR预测图：C:/xxq/output/gpr_co2_prediction.png', style='List Bullet')

# 保存文档
doc.save('C:/xxq/工作总结_7月9日.docx')
print('Word文档已生成：C:/xxq/工作总结_7月9日.docx')
